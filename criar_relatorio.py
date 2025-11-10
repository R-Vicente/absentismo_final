from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_custom(doc, text, level=1):
    """Adiciona título personalizado"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 51, 102)  # Azul escuro
    return heading

def add_paragraph_formatted(doc, text, bold=False, italic=False, size=11):
    """Adiciona parágrafo formatado"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.bold = bold
    run.italic = italic
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table_styled(doc, data, headers):
    """Adiciona tabela estilizada"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalhos
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Dados
    for row_data in data:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    return table

# Criar documento
doc = Document()

# Configurar margens
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ============================================================================
# CAPA
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ANÁLISE DE ABSENTISMO')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Call Center')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle2.add_run('Relatório Executivo para a Direção')
run.font.size = Pt(16)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Período de Análise: Janeiro 2024 – Junho 2025\n')
run.font.size = Pt(12)
run = info.add_run('3.135 Colaboradores | 1,3M Registos')
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

date = doc.add_paragraph()
date.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date.add_run('Novembro 2025')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============================================================================
# ÍNDICE
# ============================================================================
add_heading_custom(doc, 'ÍNDICE', level=1)
doc.add_paragraph()

sections_list = [
    '1. SUMÁRIO EXECUTIVO',
    '2. CONTEXTO E METODOLOGIA',
    '3. SITUAÇÃO ATUAL: INDICADORES-CHAVE',
    '4. ANÁLISE DE PADRÕES DE AUSÊNCIA',
    '5. IDENTIFICAÇÃO DE RISCOS OPERACIONAIS',
    '6. ANÁLISE COMPORTAMENTAL',
    '7. ANÁLISE TEMPORAL E TENDÊNCIAS',
    '8. SEGMENTAÇÃO E PRIORIZAÇÃO',
    '9. CONCLUSÕES E RECOMENDAÇÕES'
]

for item in sections_list:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# ============================================================================
# 1. SUMÁRIO EXECUTIVO
# ============================================================================
add_heading_custom(doc, '1. SUMÁRIO EXECUTIVO', level=1)
doc.add_paragraph()

add_paragraph_formatted(doc, 
    'Este relatório apresenta uma análise abrangente do absentismo no Call Center, '
    'baseada em 18 meses de dados operacionais (janeiro de 2024 a junho de 2025), '
    'abrangendo 3.135 colaboradores e 1,3 milhões de registos.')

doc.add_paragraph()

add_heading_custom(doc, 'Principais Conclusões', level=2)

conclusions = [
    ('Taxa de Absentismo Global:', '3,30%, com tendência de melhoria significativa entre 2024 e 2025.'),
    ('Redução Ano-a-Ano:', 'Diminuição de 36,5% na taxa de absentismo (de 3,73% em jan-jun 2024 para 2,37% em jan-jun 2025).'),
    ('Concentração do Problema:', '15 operações representam 92% do absentismo total, permitindo intervenções focalizadas.'),
    ('Padrões Disruptivos:', '74,5% dos episódios de ausência são de curta duração (até 3 dias), indicando padrões comportamentais.'),
    ('População em Risco:', '493 colaboradores ativos (42% da base ativa em 2025) apresentam episódios curtos recorrentes.'),
    ('Casos Críticos:', '42 colaboradores (8,5%) encontram-se em risco alto/severo/crítico, requerendo intervenção imediata.')
]

for label, text in conclusions:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)

doc.add_paragraph()

add_heading_custom(doc, 'Impacto Operacional', level=2)

add_paragraph_formatted(doc,
    'O absentismo representa 24.511 dias de trabalho perdidos no período analisado, '
    'equivalente a uma média de 7,8 dias por colaborador. A taxa de atrasos situa-se '
    'em 5,31%, afetando 72% dos colaboradores.')

doc.add_page_break()

# ============================================================================
# 2. CONTEXTO E METODOLOGIA
# ============================================================================
add_heading_custom(doc, '2. CONTEXTO E METODOLOGIA', level=1)
doc.add_paragraph()

add_heading_custom(doc, '2.1 Âmbito da Análise', level=2)

add_paragraph_formatted(doc,
    'A análise foi conduzida sobre um conjunto de dados operacionais do Call Center, '
    'estruturado da seguinte forma:')

doc.add_paragraph()

data = [
    ['Período', 'Janeiro 2024 – Junho 2025 (18 meses)'],
    ['Colaboradores', '3.135 colaboradores únicos'],
    ['Registos', '1.325.097 registos de presença/ausência'],
    ['Dias-colaborador', '761.244 dias úteis analisados']
]
add_table_styled(doc, data, ['Dimensão', 'Valor'])

doc.add_paragraph()

add_heading_custom(doc, '2.2 Abordagem Metodológica', level=2)

add_paragraph_formatted(doc,
    'A metodologia aplicada seguiu padrões internacionais de análise de absentismo, '
    'com foco em métricas acionáveis:')

doc.add_paragraph()

methods = [
    'Limpeza e normalização de dados, incluindo resolução de 48 incompatibilidades identificadas.',
    'Separação de dados em duas hierarquias independentes (absentismo e atrasos) para evitar dupla contagem.',
    'Cálculo de taxas normalizadas ajustadas pela base de trabalho efetivo em cada período.',
    'Análise de episódios de ausência (spells) para distinguir padrões de frequência vs. duração.',
    'Aplicação do Bradford Factor para identificar ausências disruptivas de curta duração.',
    'Segmentação por operação, categoria profissional e senioridade.',
    'Análise estatística de controlo (U-Charts) para deteção de variações anormais.'
]

for method in methods:
    p = doc.add_paragraph(method, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_heading_custom(doc, '2.3 Definições-Chave', level=2)

add_paragraph_formatted(doc,
    'Para garantir clareza interpretativa, as seguintes definições foram aplicadas:')

doc.add_paragraph()

definitions = [
    ('Taxa de Absentismo:', 'Percentagem de dias de falta (justificada + injustificada) sobre o total de dias úteis esperados (trabalho pago + faltas). Exclui férias, licenças e formações planeadas.'),
    ('Spell (Episódio):', 'Sequência contínua de dias de ausência. Um colaborador com 3 dias consecutivos de falta tem 1 spell; se faltar 3 dias separados, tem 3 spells.'),
    ('Short-term Spell:', 'Episódio de ausência até 3 dias. Estatisticamente associado a padrões comportamentais disruptivos.'),
    ('Bradford Factor:', 'Métrica que penaliza a frequência de ausências: Bradford = S² × D, onde S é o número de episódios e D o total de dias ausentes.')
]

for term, definition in definitions:
    p = doc.add_paragraph()
    run = p.add_run(term + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(definition)
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# 3. SITUAÇÃO ATUAL: INDICADORES-CHAVE
# ============================================================================
add_heading_custom(doc, '3. SITUAÇÃO ATUAL: INDICADORES-CHAVE', level=1)
doc.add_paragraph()

add_heading_custom(doc, '3.1 Visão Global', level=2)

add_paragraph_formatted(doc,
    'A situação atual do absentismo no Call Center apresenta-se da seguinte forma:')

doc.add_paragraph()

kpis = [
    ['Taxa de Absentismo', '3,30%'],
    ['Taxa de Atrasos', '5,31%'],
    ['Lost Time Rate', '7,8 dias/colaborador'],
    ['Frequency Rate', '2,43 episódios/colaborador'],
    ['Duração Média de Episódio', '3,2 dias'],
    ['Colaboradores sem Ausências', '36,2%']
]
add_table_styled(doc, kpis, ['Indicador', 'Valor'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'A taxa de absentismo de 3,30% situa-se dentro de parâmetros aceitáveis para operações '
    'de Call Center (benchmark típico: 3-5%), mas apresenta oportunidades de melhoria, '
    'particularmente considerando a tendência positiva já observada.')

doc.add_paragraph()

add_heading_custom(doc, '3.2 Distribuição por Tipo de Ausência', level=2)

add_paragraph_formatted(doc,
    'As ausências distribuem-se da seguinte forma:')

doc.add_paragraph()

absences = [
    ['Ausência Médica', '16.102 dias', '65,7%'],
    ['Falta Injustificada', '5.576 dias', '22,8%'],
    ['Assistência Familiar', '1.102 dias', '4,5%'],
    ['Óbito', '1.027 dias', '4,2%'],
    ['Casamento', '407 dias', '1,7%'],
    ['Outras', '297 dias', '1,2%']
]
add_table_styled(doc, absences, ['Tipo', 'Volume', '% Total'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'As ausências médicas representam dois terços do absentismo total. No entanto, '
    'é relevante notar que 65,7% destas ausências médicas são de curta duração (até 3 dias), '
    'o que pode indicar situações de menor gravidade ou, em alguns casos, utilização '
    'inadequada do mecanismo.')

doc.add_paragraph()

add_heading_custom(doc, '3.3 Variação por Dia da Semana', level=2)

add_paragraph_formatted(doc,
    'A taxa de absentismo apresenta variação significativa ao longo da semana:')

doc.add_paragraph()

weekdays = [
    ['Segunda-feira', '2,83%'],
    ['Terça-feira', '3,03%'],
    ['Quarta-feira', '3,09%'],
    ['Quinta-feira', '3,07%'],
    ['Sexta-feira', '3,14%'],
    ['Sábado', '4,77%'],
    ['Domingo', '4,29%']
]
add_table_styled(doc, weekdays, ['Dia da Semana', 'Taxa de Absentismo'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Observa-se um aumento substancial da taxa de absentismo ao fim de semana '
    '(sábados e domingos apresentam taxas 50-70% superiores aos dias úteis). '
    'Este padrão é comum em operações de Call Center, relacionando-se com menor '
    'apetência para trabalho em horários não convencionais e possíveis desafios '
    'de gestão de turnos.')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: Taxa de Absentismo por Dia da Semana]', italic=True)

doc.add_page_break()

# ============================================================================
# 4. ANÁLISE DE PADRÕES DE AUSÊNCIA
# ============================================================================
add_heading_custom(doc, '4. ANÁLISE DE PADRÕES DE AUSÊNCIA', level=1)
doc.add_paragraph()

add_heading_custom(doc, '4.1 Episódios de Ausência: Frequência vs. Duração', level=2)

add_paragraph_formatted(doc,
    'Foram identificados 7.600 episódios (spells) de ausência no período analisado. '
    'A distribuição por duração revela padrões importantes:')

doc.add_paragraph()

spells_dist = [
    ['1 dia', '2.953 episódios', '38,9%'],
    ['2-3 dias', '2.706 episódios', '35,6%'],
    ['4-7 dias', '1.661 episódios', '21,9%'],
    ['8-14 dias', '156 episódios', '2,1%'],
    ['Mais de 14 dias', '124 episódios', '1,6%']
]
add_table_styled(doc, spells_dist, ['Duração', 'Volume', '% do Total'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Este padrão é particularmente relevante: 74,5% dos episódios são de curta duração '
    '(até 3 dias), o que sugere que o problema principal não é a severidade médica, '
    'mas sim a frequência de ausências pontuais.')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: Distribuição de Episódios por Duração]', italic=True)

doc.add_paragraph()

add_heading_custom(doc, '4.2 Diferença de Impacto Operacional', level=2)

add_paragraph_formatted(doc,
    'Embora os episódios curtos sejam maioria em número, a análise de impacto '
    'total (dias perdidos) revela um equilíbrio:')

doc.add_paragraph()

impact = [
    ['Episódios Short-term (≤3 dias)', '5.659 episódios', '9.400 dias perdidos'],
    ['Episódios Long-term (>14 dias)', '124 episódios', '4.850 dias perdidos']
]
add_table_styled(doc, impact, ['Categoria', 'Número', 'Dias Perdidos'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Conclusão: Apesar de representarem apenas 1,6% dos episódios, as ausências longas '
    '(superiores a 14 dias) consomem cerca de 20% do total de dias perdidos. '
    'Estes casos, tipicamente associados a doenças graves ou recuperações prolongadas, '
    'requerem abordagem diferenciada, com foco em acompanhamento de saúde ocupacional.')

doc.add_paragraph()

add_heading_custom(doc, '4.3 Sazonalidade e Padrões Temporais', level=2)

add_paragraph_formatted(doc,
    'A análise temporal revela padrões sazonais consistentes:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Padrão de início de mês: ')
run.bold = True
run = p.add_run('Taxa de início de episódios nos primeiros 3 dias do mês é 15% superior à média, '
                'possivelmente relacionado com ciclos de pagamento e gestão de finanças pessoais.')

p = doc.add_paragraph()
run = p.add_run('Padrão de fim de mês: ')
run.bold = True
run = p.add_run('Taxa de início nos últimos dias do mês (28-31) é 12% superior à média.')

p = doc.add_paragraph()
run = p.add_run('Segunda-feira: ')
run.bold = True
run = p.add_run('22% dos episódios short-term iniciam à segunda-feira (esperado: 20% numa '
                'distribuição uniforme).')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: Taxa de Início de Episódios por Período do Mês]', italic=True)

doc.add_page_break()

# ============================================================================
# 5. IDENTIFICAÇÃO DE RISCOS OPERACIONAIS
# ============================================================================
add_heading_custom(doc, '5. IDENTIFICAÇÃO DE RISCOS OPERACIONAIS', level=1)
doc.add_paragraph()

add_heading_custom(doc, '5.1 Bradford Factor: Deteção de Padrões Disruptivos', level=2)

add_paragraph_formatted(doc,
    'O Bradford Factor é uma métrica internacional utilizada para identificar ausências '
    'de alta disrupção operacional. A fórmula penaliza a frequência de episódios: '
    'um colaborador com 10 episódios de 1 dia (Bradford = 1.000) é considerado 100 vezes '
    'mais disruptivo que um colaborador com 1 episódio de 10 dias (Bradford = 10).')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'A análise foi aplicada exclusivamente a colaboradores ativos em 2025, '
    'considerando apenas episódios short-term (até 3 dias), gerando o "Bradford Disruptivo":')

doc.add_paragraph()

bradford = [
    ['Aceitável (< 50)', '383 colaboradores', '77,7%'],
    ['Monitorizar (50-100)', '41 colaboradores', '8,3%'],
    ['Atenção (100-200)', '27 colaboradores', '5,5%'],
    ['Alto Risco (200-400)', '19 colaboradores', '3,9%'],
    ['Risco Severo (400-600)', '10 colaboradores', '2,0%'],
    ['Risco Crítico (> 600)', '13 colaboradores', '2,6%']
]
add_table_styled(doc, bradford, ['Nível de Risco', 'Volume', '% do Total'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Análise: 493 colaboradores ativos (41,6% da base ativa em 2025) apresentam episódios '
    'short-term recorrentes. Destes, 69 colaboradores (14%) encontram-se acima do limiar '
    'de 100 pontos, indicando necessidade de intervenção. Os 42 colaboradores em risco '
    'alto/severo/crítico (8,5%) devem ser priorizados para ação imediata.')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: Distribuição por Nível de Risco Bradford]', italic=True)

doc.add_paragraph()

add_heading_custom(doc, '5.2 Segmentação por Operação', level=2)

add_paragraph_formatted(doc,
    'Algumas operações apresentam níveis de Bradford Disruptivo significativamente '
    'superiores à média (entre operações com mínimo de 20 colaboradores):')

doc.add_paragraph()

ops_bradford = [
    ['UNITEL NET CASA 5G (Outsourcing)', '200,2', '45 colaboradores'],
    ['Unitel Alpha', '97,2', '114 colaboradores'],
    ['Unitel Money - Brigadistas (Outsourcing)', '96,9', '77 colaboradores'],
    ['TAAG', '70,9', '44 colaboradores']
]
add_table_styled(doc, ops_bradford, ['Operação', 'Bradford Médio', 'N Colaboradores'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Destaque: A operação UNITEL NET CASA 5G (Outsourcing) encontra-se em situação crítica, '
    'com Bradford médio de 200,2, o dobro do limiar de atenção e o triplo da média geral.')

doc.add_paragraph()

add_heading_custom(doc, '5.3 Segmentação por Categoria Profissional', level=2)

add_paragraph_formatted(doc,
    'A análise por função revela disparidades significativas:')

doc.add_paragraph()

cats_bradford = [
    ['Comercial', '144,8', '79 colaboradores'],
    ['Brigadista', '93,6', '81 colaboradores'],
    ['Assistente de Contact Center', '50,7', '321 colaboradores']
]
add_table_styled(doc, cats_bradford, ['Categoria', 'Bradford Médio', 'N Colaboradores'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Observação: A função de Comercial apresenta Bradford médio 3 vezes superior à dos '
    'Assistentes de Contact Center, apesar de representar um grupo menor. '
    'Esta categoria merece atenção específica.')

doc.add_page_break()

# ============================================================================
# 6. ANÁLISE COMPORTAMENTAL
# ============================================================================
add_heading_custom(doc, '6. ANÁLISE COMPORTAMENTAL', level=1)
doc.add_paragraph()

add_paragraph_formatted(doc,
    'Para além da análise quantitativa, foram aplicadas técnicas de deteção de padrões '
    'comportamentais suspeitos, com base em anomalias estatísticas.')

doc.add_paragraph()

add_heading_custom(doc, '6.1 Padrão Segunda-Feira/Sexta-Feira', level=2)

add_paragraph_formatted(doc,
    'Análise estatística identifica concentração anormal de episódios nas extremidades '
    'da semana:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Baseline esperado: ')
run.bold = True
run = p.add_run('20% dos episódios deveriam iniciar em cada dia útil (distribuição uniforme).')

p = doc.add_paragraph()
run = p.add_run('Observado: ')
run.bold = True
run = p.add_run('22% dos episódios short-term iniciam à segunda-feira.')

p = doc.add_paragraph()
run = p.add_run('Colaboradores flagged: ')
run.bold = True
run = p.add_run('Identificados colaboradores com mais de 50% dos episódios iniciando '
                'à segunda ou terminando à sexta (mínimo 5 episódios para significância estatística).')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Este padrão pode indicar extensão não autorizada de fins de semana. '
    'Recomenda-se análise individual dos casos flagged.')

doc.add_paragraph()

add_heading_custom(doc, '6.2 Padrão de Ponte (Adjacente a Feriados)', level=2)

add_paragraph_formatted(doc,
    'A análise de 23 feriados nacionais no período revelou:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('14,7% ')
run.bold = True
run = p.add_run('dos episódios short-term iniciam adjacentes a feriados (±1 dia).')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Esta taxa está ligeiramente acima do esperado, mas não foram identificados '
    'colaboradores individuais com padrão consistente de "fazer ponte" (critério: >40% '
    'dos episódios adjacentes a feriados).')

doc.add_paragraph()

add_heading_custom(doc, '6.3 Baixas Médicas de Curta Duração', level=2)

add_paragraph_formatted(doc,
    'Das 16.102 ausências médicas registadas, uma proporção substancial é de curta duração:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Baixas médicas ≤ 3 dias: ')
run.bold = True
run = p.add_run('10.580 episódios (65,7% das ausências médicas).')

p = doc.add_paragraph()
run = p.add_run('Concentração em segundas-feiras: ')
run.bold = True
run = p.add_run('23,5% das baixas médicas curtas iniciam à segunda.')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Embora baixas médicas curtas sejam legítimas na maioria dos casos (gripe, '
    'mal-estar pontual), a concentração ligeiramente acima do esperado em segundas-feiras '
    'sugere que uma pequena fração pode representar utilização inadequada. '
    'Foram identificados colaboradores com padrão recorrente (≥3 baixas médicas curtas, '
    '>60% em segundas) para eventual acompanhamento.')

doc.add_paragraph()

add_heading_custom(doc, '6.4 Síntese: Colaboradores com Múltiplos Flags', level=2)

add_paragraph_formatted(doc,
    'A combinação de múltiplos indicadores comportamentais permite priorização:')

doc.add_paragraph()

flags = [
    ['Sem flags suspeitos', 'Maioria dos colaboradores'],
    ['1 flag', 'Atenção ligeira'],
    ['2 flags', 'Alta suspeita - revisão recomendada'],
    ['3 flags', 'Muito suspeito - investigação prioritária']
]
add_table_styled(doc, flags, ['Número de Flags', 'Interpretação'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Nota metodológica: Estes indicadores não constituem prova de má conduta, '
    'mas sim ferramentas estatísticas de triagem. Qualquer ação individual deve ser '
    'precedida de análise contextual e conversa com o colaborador e sua chefia direta.')

doc.add_page_break()

# ============================================================================
# 7. ANÁLISE TEMPORAL E TENDÊNCIAS
# ============================================================================
add_heading_custom(doc, '7. ANÁLISE TEMPORAL E TENDÊNCIAS', level=1)
doc.add_paragraph()

add_heading_custom(doc, '7.1 Comparação Ano-a-Ano: 2024 vs. 2025', level=2)

add_paragraph_formatted(doc,
    'Para garantir comparabilidade, foram analisados apenas os primeiros seis meses '
    'de cada ano (janeiro a junho):')

doc.add_paragraph()

comparison = [
    ['2024 (Jan-Jun)', '3,73%', '8.934 dias', '2.610 episódios'],
    ['2025 (Jan-Jun)', '2,37%', '6.026 dias', '2.185 episódios'],
    ['Variação', '-36,5%', '-32,5%', '-16,3%']
]
add_table_styled(doc, comparison, ['Período', 'Taxa Absentismo', 'Dias Perdidos', 'N Episódios'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Conclusão principal: Regista-se uma melhoria substancial e consistente em todos '
    'os indicadores de absentismo entre 2024 e 2025. A redução de 36,5% na taxa de '
    'absentismo é particularmente significativa e sugere que medidas implementadas '
    'ao longo de 2024 (se aplicável) estão a produzir resultados positivos.')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: Evolução Mensal da Taxa de Absentismo - 2024 vs 2025]', italic=True)

doc.add_paragraph()

add_heading_custom(doc, '7.2 Tendência Intra-Ano', level=2)

add_paragraph_formatted(doc,
    'A análise mês a mês revela consistência na melhoria:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Todos os meses de 2025 ')
run = p.add_run('apresentam taxas inferiores aos meses equivalentes de 2024.')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Maior redução observada: ')
run.bold = True
run = p.add_run('Março (de 4,1% em 2024 para 2,3% em 2025).')
run.font.size = Pt(11)

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Esta consistência reforça a hipótese de que a melhoria não resulta de fatores '
    'sazonais ou pontuais, mas sim de mudança estrutural.')

doc.add_paragraph()

add_heading_custom(doc, '7.3 Duração de Episódios: Evolução', level=2)

add_paragraph_formatted(doc,
    'Para além da frequência, também a duração média dos episódios diminuiu:')

doc.add_paragraph()

duration = [
    ['2024 (Jan-Jun)', '3,7 dias'],
    ['2025 (Jan-Jun)', '2,7 dias'],
    ['Redução', '27%']
]
add_table_styled(doc, duration, ['Período', 'Duração Média'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Interpretação: Os colaboradores não apenas faltam menos vezes, mas quando faltam, '
    'os episódios são mais curtos. Esta dupla melhoria (frequência + duração) é indicador '
    'robusto de progresso.')

doc.add_paragraph()

add_heading_custom(doc, '7.4 Análise de Controlo Estatístico (U-Chart)', level=2)

add_paragraph_formatted(doc,
    'Para monitorizar estabilidade do processo, foi aplicado U-Chart (control chart para '
    'dados de contagem) às semanas de 2025:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Semanas fora de controlo estatístico: ')
run.bold = True
run = p.add_run('Identificadas semanas com variação anormal (acima do limite superior '
                'ou abaixo do limite inferior), requerendo investigação de causas especiais.')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Recomendação: Quando uma semana excede os limites de controlo, deve ser realizada '
    'análise de root cause (exemplo: surto de gripe, evento específico na operação, '
    'problema de transporte coletivo).')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: U-Chart - Taxa Semanal de Faltas em 2025]', italic=True)

doc.add_page_break()

# ============================================================================
# 8. SEGMENTAÇÃO E PRIORIZAÇÃO
# ============================================================================
add_heading_custom(doc, '8. SEGMENTAÇÃO E PRIORIZAÇÃO', level=1)
doc.add_paragraph()

add_heading_custom(doc, '8.1 Princípio de Pareto: Concentração do Problema', level=2)

add_paragraph_formatted(doc,
    'A análise de contribuição revela concentração significativa do absentismo em '
    'poucos segmentos:')

doc.add_paragraph()

pareto_ops = [
    ['Top 5 operações', '76% do absentismo total'],
    ['Top 10 operações', '88% do absentismo total'],
    ['Top 15 operações', '92% do absentismo total']
]
add_table_styled(doc, pareto_ops, ['Segmento', 'Contribuição Acumulada'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Implicação estratégica: Focar esforços de melhoria em 15 operações permite '
    'impactar 92% do problema. Esta concentração facilita intervenções dirigidas '
    'e medição de eficácia.')

doc.add_paragraph()

add_heading_custom(doc, '8.2 Matriz Taxa vs. Contribuição', level=2)

add_paragraph_formatted(doc,
    'Para priorização eficaz, é útil distinguir duas dimensões:')

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run('Taxa de absentismo: ')
run.bold = True
run = p.add_run('Percentagem de ausências no segmento (indicador de severidade do problema).')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('Contribuição absoluta: ')
run.bold = True
run = p.add_run('Número de dias perdidos (indicador de impacto no negócio).')
run.font.size = Pt(11)

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Operações prioritárias são aquelas que combinam taxa elevada E contribuição elevada:')

doc.add_paragraph()

priority_ops = [
    ['Unitel Alpha', '4,81%', '8.716 dias (35,6% do total)'],
    ['Unitel Money - Brigadistas', '4,59%', '1.706 dias (7,0%)'],
    ['ZAP Fibra - BO/SWAT', '10,09%', '307 dias (1,3%)']
]
add_table_styled(doc, priority_ops, ['Operação', 'Taxa', 'Contribuição'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Destaque: Unitel Alpha combina taxa acima da média com volume absoluto dominante, '
    'representando sozinha mais de um terço do absentismo total da empresa.')

doc.add_paragraph()

add_paragraph_formatted(doc, '[Gráfico: Matriz Taxa vs Contribuição - Operações]', italic=True)

doc.add_paragraph()

add_heading_custom(doc, '8.3 Clustering: Perfis de Comportamento', level=2)

add_paragraph_formatted(doc,
    'Aplicando técnicas de clustering (K-Means) aos colaboradores com episódios short-term, '
    'foram identificados 3 perfis distintos:')

doc.add_paragraph()

clusters = [
    ['Cluster 1: Baixo Risco', '~350 colaboradores', 'Bradford < 50, poucos episódios'],
    ['Cluster 2: Risco Moderado', '~70 colaboradores', 'Bradford 50-150, frequência média'],
    ['Cluster 3: Alto Risco', '~70 colaboradores', 'Bradford > 150, padrão disruptivo']
]
add_table_styled(doc, clusters, ['Perfil', 'Volume', 'Características'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Esta segmentação permite estratégias diferenciadas: o Cluster 3 requer intervenção '
    'individual e acompanhamento próximo; o Cluster 2 beneficia de ações de sensibilização; '
    'o Cluster 1 necessita apenas monitorização de rotina.')

doc.add_paragraph()

add_heading_custom(doc, '8.4 Análise por Senioridade (Cohorts)', level=2)

add_paragraph_formatted(doc,
    'A segmentação por tempo de casa revela diferenças importantes:')

doc.add_paragraph()

cohorts = [
    ['< 1 ano', '2,95%'],
    ['1-2 anos', '2,48%'],
    ['2-3 anos', '2,31%'],
    ['3-5 anos', '2,18%'],
    ['> 5 anos', '2,85%']
]
add_table_styled(doc, cohorts, ['Senioridade', 'Taxa de Absentismo'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Observações: Colaboradores muito novos (<1 ano) e muito antigos (>5 anos) apresentam '
    'taxas ligeiramente superiores. Para novos colaboradores, pode relacionar-se com período '
    'de adaptação; para veteranos, possível burnout ou questões de saúde acumuladas. '
    'A faixa 3-5 anos apresenta o melhor desempenho.')

doc.add_page_break()

# ============================================================================
# 9. CONCLUSÕES E RECOMENDAÇÕES
# ============================================================================
add_heading_custom(doc, '9. CONCLUSÕES E RECOMENDAÇÕES', level=1)
doc.add_paragraph()

add_heading_custom(doc, '9.1 Síntese de Conclusões', level=2)

conclusions_final = [
    ('Tendência positiva:', 'A taxa de absentismo reduziu 36,5% entre 2024 e 2025, sinalizando eficácia de eventuais medidas já implementadas.'),
    ('Problema concentrado:', '15 operações representam 92% do absentismo, permitindo intervenções focalizadas com alto retorno.'),
    ('Padrão predominante:', '74,5% dos episódios são short-term (≤3 dias), indicando que o desafio principal é comportamental, não médico.'),
    ('População em risco identificada:', '493 colaboradores com episódios recorrentes, dos quais 42 em risco crítico.'),
    ('Variação semanal:', 'Fim de semana apresenta taxa 50-70% superior, sugerindo desafios específicos de gestão de turnos.'),
    ('Oportunidade de melhoria:', 'Apesar da tendência positiva, subsistem segmentos e colaboradores com padrões disruptivos corrigíveis.')
]

for i, (label, text) in enumerate(conclusions_final, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {label} ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)

doc.add_paragraph()

add_heading_custom(doc, '9.2 Recomendações Estratégicas', level=2)

add_paragraph_formatted(doc,
    'Com base na análise conduzida, recomendam-se as seguintes linhas de ação:')

doc.add_paragraph()

add_heading_custom(doc, 'A. Intervenções Prioritárias (Curto Prazo - 0-3 meses)', level=3)

short_term = [
    'Revisão individual dos 42 colaboradores em risco crítico (Bradford > 200), com conversas estruturadas envolvendo RH e chefia direta.',
    'Análise de root cause na operação UNITEL NET CASA 5G (Outsourcing), cuja taxa é 2x superior à média.',
    'Implementação de política clara sobre baixas médicas short-term, incluindo requisitos de certificação médica a partir do 2º episódio num período de 30 dias.',
    'Comunicação de política de absentismo atualizada, enfatizando consequências progressivas para padrões recorrentes sem justificação adequada.'
]

for i, rec in enumerate(short_term, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(rec)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_heading_custom(doc, 'B. Melhorias Estruturais (Médio Prazo - 3-6 meses)', level=3)

medium_term = [
    'Implementação de dashboard de monitorização em tempo real (mensal) do Bradford Factor por operação e colaborador.',
    'Programa de reconhecimento para colaboradores com zero ausências ou melhoria significativa face a períodos anteriores.',
    'Revisão da gestão de turnos de fim de semana, incluindo análise de incentivos, rotatividade e condições específicas.',
    'Formação para chefias de equipa em gestão de absentismo, incluindo conversas difíceis e identificação de padrões.',
    'Análise de causas raiz nas operações Top 5 por contribuição, com planos de ação específicos para cada uma.'
]

for i, rec in enumerate(medium_term, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(rec)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_heading_custom(doc, 'C. Iniciativas de Longo Prazo (6-12 meses)', level=3)

long_term = [
    'Programa de bem-estar e saúde ocupacional, com foco preventivo em doenças recorrentes.',
    'Análise de correlação entre absentismo e outras métricas (performance, satisfação, turnover) para identificar padrões preditivos.',
    'Implementação de sistema de alerta automático para chefias quando colaborador atinge threshold de Bradford.',
    'Revisão de processos de recrutamento e onboarding para reduzir absentismo em colaboradores <1 ano.',
    'Benchmark externo com outras operações de Call Center para validar metas de taxa de absentismo.'
]

for i, rec in enumerate(long_term, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(rec)
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_heading_custom(doc, '9.3 Métricas de Acompanhamento', level=2)

add_paragraph_formatted(doc,
    'Para medir o progresso das iniciativas, recomenda-se monitorização trimestral dos seguintes indicadores:')

doc.add_paragraph()

metrics = [
    'Taxa de absentismo global',
    'Número de colaboradores com Bradford > 200',
    'Taxa de absentismo nas Top 5 operações',
    'Percentagem de colaboradores com zero ausências',
    'Número de semanas fora de controlo estatístico (U-Chart)',
    'Taxa de absentismo em fins de semana vs. dias úteis'
]

for metric in metrics:
    p = doc.add_paragraph(metric, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_heading_custom(doc, '9.4 Considerações Finais', level=2)

add_paragraph_formatted(doc,
    'A análise demonstra que o absentismo no Call Center, embora dentro de parâmetros '
    'aceitáveis e com tendência positiva, apresenta oportunidades claras de melhoria. '
    'A concentração do problema em segmentos específicos e a natureza predominantemente '
    'comportamental (episódios curtos e frequentes) sugerem que intervenções dirigidas '
    'podem produzir resultados significativos.')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Recomenda-se abordagem equilibrada, combinando:')

doc.add_paragraph()

approaches = [
    'Ações corretivas para casos críticos identificados',
    'Melhorias estruturais em processos e políticas',
    'Iniciativas preventivas de bem-estar e engagement',
    'Monitorização contínua com dashboards e alertas automáticos'
]

for approach in approaches:
    p = doc.add_paragraph(approach, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_paragraph()

add_paragraph_formatted(doc,
    'A tendência de melhoria observada entre 2024 e 2025 é encorajadora e deve ser '
    'preservada e acelerada através das recomendações apresentadas.')

doc.add_page_break()

# ============================================================================
# ANEXOS
# ============================================================================
add_heading_custom(doc, 'ANEXOS', level=1)
doc.add_paragraph()

add_heading_custom(doc, 'Anexo A: Glossário Técnico', level=2)

glossary = [
    ('Absentismo:', 'Ausência não planeada do local de trabalho durante o horário de trabalho esperado.'),
    ('Bradford Factor:', 'Métrica de disrupção operacional calculada como S² × D, onde S é o número de episódios e D o total de dias ausentes.'),
    ('Spell (Episódio):', 'Sequência contínua de dias de ausência, independentemente da duração.'),
    ('Short-term Spell:', 'Episódio de ausência com duração até 3 dias.'),
    ('Long-term Spell:', 'Episódio de ausência com duração superior a 14 dias.'),
    ('Taxa de Absentismo:', 'Percentagem de dias de falta sobre o total de dias úteis esperados (Faltas / (Trabalho Pago + Faltas) × 100).'),
    ('Lost Time Rate:', 'Média de dias perdidos por colaborador no período.'),
    ('Frequency Rate:', 'Média de episódios de ausência por colaborador.'),
    ('U-Chart:', 'Gráfico de controlo estatístico para monitorizar taxas de eventos raros ao longo do tempo.'),
    ('Cluster:', 'Grupo de colaboradores com perfil similar de absentismo, identificado por análise estatística.')
]

for term, definition in glossary:
    p = doc.add_paragraph()
    run = p.add_run(term + ' ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(definition)
    run.font.size = Pt(10)

doc.add_paragraph()

add_heading_custom(doc, 'Anexo B: Metodologia do Bradford Factor', level=2)

add_paragraph_formatted(doc,
    'O Bradford Factor foi desenvolvido na década de 1980 na Universidade de Bradford (Reino Unido) '
    'e é amplamente utilizado em operações com elevada sensibilidade a ausências pontuais, '
    'como Call Centers, hospitais e linhas de produção.')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Fórmula: Bradford = S² × D', bold=True)

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Onde:')

p = doc.add_paragraph()
run = p.add_run('S = ')
run.bold = True
run = p.add_run('Número de spells (episódios) num período definido')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('D = ')
run.bold = True
run = p.add_run('Total de dias ausentes no mesmo período')
run.font.size = Pt(11)

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Exemplo ilustrativo:')

doc.add_paragraph()

examples = [
    ['Colaborador A', '1 episódio de 10 dias', 'Bradford = 1² × 10 = 10'],
    ['Colaborador B', '5 episódios de 2 dias', 'Bradford = 5² × 10 = 250'],
    ['Colaborador C', '10 episódios de 1 dia', 'Bradford = 10² × 10 = 1.000']
]
add_table_styled(doc, examples, ['Perfil', 'Padrão', 'Bradford'])

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Interpretação: Apesar de todos terem 10 dias de ausência, o Colaborador C '
    '(10 episódios de 1 dia) tem Bradford 100x superior ao Colaborador A (1 episódio de 10 dias), '
    'refletindo a maior disrupção operacional causada por ausências frequentes e imprevisíveis.')

doc.add_paragraph()

add_paragraph_formatted(doc,
    'Thresholds aplicados neste relatório (baseados em Call Centre Helper UK):')

doc.add_paragraph()

thresholds = [
    ['< 50', 'Aceitável', 'Monitorização normal'],
    ['50-100', 'Monitorizar', 'Prestar atenção'],
    ['100-200', 'Atenção', 'Conversa informal'],
    ['200-400', 'Alto Risco', 'Reunião formal'],
    ['400-600', 'Risco Severo', 'Aviso formal'],
    ['> 600', 'Risco Crítico', 'Ação disciplinar']
]
add_table_styled(doc, thresholds, ['Score', 'Categoria', 'Ação Recomendada'])

doc.add_paragraph()

add_heading_custom(doc, 'Anexo C: Notas Metodológicas', level=2)

add_paragraph_formatted(doc,
    'Decisões metodológicas críticas aplicadas nesta análise:')

doc.add_paragraph()

notes = [
    'Separação de hierarquias: Dados de absentismo e atrasos foram tratados separadamente para evitar dupla contagem (um colaborador pode trabalhar E ter atraso no mesmo dia).',
    'Normalização temporal: Todas as análises por dia da semana ou período do mês utilizam taxas normalizadas pela base de trabalho, não valores absolutos.',
    'Bradford Disruptivo: Aplicado apenas a colaboradores ativos em 2025, considerando exclusivamente spells short-term (≤3 dias) para focar em padrões comportamentais.',
    'Filtro de amostra mínima: Rankings por operação ou categoria excluem segmentos com <20 colaboradores para evitar distorções estatísticas.',
    'Comparação anual: Comparação 2024 vs 2025 limitada a janeiro-junho de ambos os anos para garantir equivalência.',
    'Incompatibilidades: 48 dias com registos contraditórios (ex: presença + ausência médica no mesmo dia) foram excluídos após validação.'
]

for i, note in enumerate(notes, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(note)
    run.font.size = Pt(10)

# Salvar documento
doc.save('/home/user/absentismo_final/Relatorio_Absentismo_Direcao_Executiva.docx')
print('Relatório criado com sucesso!')
